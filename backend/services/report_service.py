"""
Report Service Layer

Handles business logic for report generation operations following SOLID principles:
- Single Responsibility: Only handles report-related business logic
- Dependency Inversion: Depends on Session abstraction
- Open/Closed: Easy to extend with new report types
"""

from sqlalchemy.orm import Session
from typing import List, Dict, Any, Optional
from datetime import datetime
from pathlib import Path
import json
from collections import defaultdict

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.database import ProjectDB, UserStoryDB, TestCaseDB, BugReportDB, TestExecutionDB
from backend.models import UserStory, TestCase, BugReport, BugSeverity, BugPriority, BugStatus, BugType
from backend.generators import TestPlanGenerator
from backend.generators.bug_report_generator import BugReportGenerator
from backend.config import settings


class ReportService:
    """Service class for report generation business logic"""

    def __init__(self, db: Session):
        """Initialize service with database session"""
        self.db = db

    def generate_test_plan(
        self,
        project_id: str,
        format: str = "both"
    ) -> Dict[str, Any]:
        """
        Generate test plan document for a specific project

        Args:
            project_id: Project ID to generate test plan for
            format: Format - "pdf", "docx", or "both"

        Returns:
            Dictionary with generated files

        Raises:
            ValueError: If project not found
        """
        # Validate project exists
        project = self.db.query(ProjectDB).filter(ProjectDB.id == project_id).first()
        if not project:
            raise ValueError(f"Project {project_id} not found")

        # Get user stories and test cases for this project only
        user_stories_db = self.db.query(UserStoryDB).filter(UserStoryDB.project_id == project_id).all()
        test_cases_db = self.db.query(TestCaseDB).filter(TestCaseDB.project_id == project_id).all()

        # Convert to models (simplified)
        user_stories = [
            UserStory(
                id=s.id,
                title=s.title,
                description=s.description,
                priority=s.priority,
                status=s.status
            )
            for s in user_stories_db
        ]

        test_cases = [
            TestCase(
                id=tc.id,
                title=tc.title,
                description=tc.description,
                user_story_id=tc.user_story_id,
                test_type=tc.test_type,
                priority=tc.priority,
                status=tc.status
            )
            for tc in test_cases_db
        ]

        # Generate test plan
        settings.ensure_directories()
        test_plan_gen = TestPlanGenerator()
        files = test_plan_gen.generate_test_plan(
            user_stories=user_stories,
            test_cases=test_cases,
            output_dir=settings.output_dir,
            project_name=project.name,
            format=format
        )

        return {
            "message": "Test plan generated successfully",
            "files": files
        }

    def generate_bug_summary_report(self, project_id: str) -> str:
        """
        Generate Bug Summary Report for Dev Team
        Returns path to generated Word document

        Args:
            project_id: Project ID

        Returns:
            Path to generated file

        Raises:
            ValueError: If project not found or no bugs found
        """
        # Validate project exists
        project = self.db.query(ProjectDB).filter(ProjectDB.id == project_id).first()
        if not project:
            raise ValueError(f"Project {project_id} not found")

        # Get all bugs for the project
        bugs_db = self.db.query(BugReportDB).filter(BugReportDB.project_id == project_id).all()

        if not bugs_db:
            raise ValueError("No bugs found for this project")

        # Convert to Pydantic models
        bugs = []
        for bug_db in bugs_db:
            bug = BugReport(
                id=bug_db.id,
                title=bug_db.title,
                description=bug_db.description,
                steps_to_reproduce=bug_db.steps_to_reproduce.split('\n') if bug_db.steps_to_reproduce else [],
                expected_behavior=bug_db.expected_behavior or "",
                actual_behavior=bug_db.actual_behavior or "",
                severity=BugSeverity(bug_db.severity),
                priority=BugPriority(bug_db.priority),
                bug_type=BugType(bug_db.bug_type),
                status=BugStatus(bug_db.status),
                environment=bug_db.environment,
                browser=bug_db.browser,
                os=bug_db.os,
                version=bug_db.version,
                user_story_id=bug_db.user_story_id,
                test_case_id=bug_db.test_case_id,
                scenario_name=bug_db.scenario_name,
                screenshots=json.loads(bug_db.screenshots) if bug_db.screenshots else [],
                logs=bug_db.logs,
                notes=bug_db.notes,
                workaround=bug_db.workaround,
                root_cause=bug_db.root_cause,
                fix_description=bug_db.fix_description,
                reported_by=bug_db.reported_by or "Unknown",
                assigned_to=bug_db.assigned_to,
                verified_by=bug_db.verified_by,
                reported_date=bug_db.reported_date,
                assigned_date=bug_db.assigned_date,
                fixed_date=bug_db.fixed_date,
                verified_date=bug_db.verified_date,
                closed_date=bug_db.closed_date
            )
            bugs.append(bug)

        # Generate report
        generator = BugReportGenerator()
        output_dir = "output/reports"
        filename = f"BugSummary_{project.name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        file_path = generator.generate_bulk_report(bugs, output_dir, filename)

        return str(file_path)

    def generate_test_execution_report(self, project_id: str) -> str:
        """
        Generate Test Execution Summary Report for QA Manager
        Returns path to generated Word document

        Args:
            project_id: Project ID

        Returns:
            Path to generated file

        Raises:
            ValueError: If project not found or no test cases found
        """
        # Validate project exists
        project = self.db.query(ProjectDB).filter(ProjectDB.id == project_id).first()
        if not project:
            raise ValueError(f"Project {project_id} not found")

        # Get test cases for the project
        test_cases = self.db.query(TestCaseDB).filter(TestCaseDB.project_id == project_id).all()

        if not test_cases:
            raise ValueError("No test cases found for this project")

        # Get all executions for these test cases
        test_case_ids = [tc.id for tc in test_cases]
        executions = self.db.query(TestExecutionDB).filter(
            TestExecutionDB.test_case_id.in_(test_case_ids)
        ).order_by(TestExecutionDB.execution_date.desc()).all()

        # Calculate statistics
        total_tests = len(test_cases)
        total_executions = len(executions)

        # Execution status counts
        status_counts = {
            'PASSED': 0,
            'FAILED': 0,
            'BLOCKED': 0,
            'SKIPPED': 0,
            'NOT_RUN': 0
        }

        for execution in executions:
            status_counts[execution.status] = status_counts.get(execution.status, 0) + 1

        # Test case status (latest execution)
        test_case_latest_status = {}
        test_case_info = {}
        for tc in test_cases:
            test_case_info[tc.id] = {
                'title': tc.title,
                'type': tc.test_type,
                'priority': tc.priority
            }

            latest = self.db.query(TestExecutionDB).filter(
                TestExecutionDB.test_case_id == tc.id
            ).order_by(TestExecutionDB.execution_date.desc()).first()

            if latest:
                test_case_latest_status[tc.id] = latest.status
            else:
                test_case_latest_status[tc.id] = 'NOT_RUN'

        # Group executions by Test Case and Scenario
        grouped_executions = self._group_executions_by_test_case_and_scenario(executions)

        # Create document
        doc = self._create_test_execution_document(
            project=project,
            total_tests=total_tests,
            total_executions=total_executions,
            status_counts=status_counts,
            grouped_executions=grouped_executions,
            test_case_info=test_case_info
        )

        # Save document
        output_dir = Path("output/reports")
        output_dir.mkdir(parents=True, exist_ok=True)
        filename = f"TestExecution_{project.name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = output_dir / filename

        doc.save(str(file_path))

        return str(file_path)

    def generate_consolidated_report(self) -> str:
        """
        Generate Consolidated Report for Manager
        Returns path to generated Word document

        Returns:
            Path to generated file

        Raises:
            ValueError: If no projects found
        """
        # Get all projects
        projects = self.db.query(ProjectDB).all()

        if not projects:
            raise ValueError("No projects found")

        # Collect metrics for each project
        project_metrics = self._calculate_project_metrics(projects)

        # Calculate totals and averages
        total_stories = sum(p['stories'] for p in project_metrics)
        total_test_cases = sum(p['test_cases'] for p in project_metrics)
        total_bugs = sum(p['bugs'] for p in project_metrics)
        total_coverage = sum(p['coverage'] for p in project_metrics)

        avg_coverage = total_coverage / len(projects) if projects else 0
        avg_health = sum(p['health_score'] for p in project_metrics) / len(projects) if projects else 0

        # Sort and filter projects
        projects_by_health = sorted(project_metrics, key=lambda p: p['health_score'], reverse=True)
        at_risk_projects = [p for p in project_metrics if p['risk_level'] in ['ALTO', 'MEDIO']]

        # Create document
        doc = self._create_consolidated_document(
            projects=projects,
            project_metrics=project_metrics,
            projects_by_health=projects_by_health,
            at_risk_projects=at_risk_projects,
            total_stories=total_stories,
            total_test_cases=total_test_cases,
            total_bugs=total_bugs,
            avg_coverage=avg_coverage,
            avg_health=avg_health
        )

        # Save document
        output_dir = Path("output/reports")
        output_dir.mkdir(parents=True, exist_ok=True)
        filename = f"Consolidated_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = output_dir / filename

        doc.save(str(file_path))

        return str(file_path)

    # ========== Private Helper Methods ==========

    def _group_executions_by_test_case_and_scenario(self, executions: List[TestExecutionDB]) -> Dict:
        """Group executions by test case and scenario"""
        grouped_executions = defaultdict(lambda: defaultdict(list))

        for execution in executions:
            test_case_id = execution.test_case_id

            if execution.step_results:
                try:
                    step_results = json.loads(execution.step_results)

                    # Group steps by scenario
                    scenario_steps = defaultdict(list)
                    for step in step_results:
                        scenario_name = step.get('scenario', 'Default Scenario')
                        scenario_steps[scenario_name].append(step)

                    # Create execution record for each scenario
                    for scenario_name, steps in scenario_steps.items():
                        grouped_executions[test_case_id][scenario_name].append({
                            'execution': execution,
                            'steps': steps
                        })
                except json.JSONDecodeError:
                    # Fallback if JSON parsing fails
                    grouped_executions[test_case_id]['Default Scenario'].append({
                        'execution': execution,
                        'steps': []
                    })
            else:
                # No step results, add to default scenario
                grouped_executions[test_case_id]['Default Scenario'].append({
                    'execution': execution,
                    'steps': []
                })

        return grouped_executions

    def _create_test_execution_document(
        self,
        project: ProjectDB,
        total_tests: int,
        total_executions: int,
        status_counts: Dict,
        grouped_executions: Dict,
        test_case_info: Dict
    ) -> Document:
        """Create Word document for test execution report"""
        doc = Document()

        # Title
        title = doc.add_heading(f"Test Execution Summary Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Project info
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"Project: {project.name}\n").bold = True
        info_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
        info_para.add_run(f"Total Test Cases: {total_tests}\n")
        info_para.add_run(f"Total Executions: {total_executions}")

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(f"Total Test Cases: {total_tests}")
        doc.add_paragraph(f"Total Test Executions: {total_executions}")

        # Execution Status Distribution
        doc.add_heading("Execution Status Distribution", level=2)
        status_table = doc.add_table(rows=1, cols=2)
        status_table.style = "Light Grid Accent 1"

        # Header
        header_cells = status_table.rows[0].cells
        header_cells[0].text = "Status"
        header_cells[1].text = "Count"
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True

        # Add status counts
        for status, count in sorted(status_counts.items()):
            row_cells = status_table.add_row().cells
            row_cells[0].text = status
            row_cells[1].text = str(count)

        # Pass Rate
        if total_executions > 0:
            pass_rate = (status_counts['PASSED'] / total_executions) * 100
            doc.add_paragraph()
            pass_rate_para = doc.add_paragraph()
            pass_rate_para.add_run(f"Overall Pass Rate: {pass_rate:.1f}%").bold = True
            if pass_rate >= 80:
                pass_rate_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)
            elif pass_rate >= 60:
                pass_rate_para.runs[0].font.color.rgb = RGBColor(255, 165, 0)
            else:
                pass_rate_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)

        # Detailed Execution Results
        doc.add_heading("Execution Results by Test Case and Scenario", level=1)

        if not grouped_executions:
            doc.add_paragraph("No execution results to display.")
        else:
            self._add_execution_details_to_document(doc, grouped_executions, test_case_info)

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph("Generated by Quality Mission Control System")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        return doc

    def _add_execution_details_to_document(self, doc: Document, grouped_executions: Dict, test_case_info: Dict):
        """Add execution details tables to document"""
        for test_case_id in sorted(grouped_executions.keys()):
            tc_info = test_case_info.get(test_case_id, {'title': 'Unknown', 'type': 'N/A', 'priority': 'N/A'})
            doc.add_heading(f"Test Case: {test_case_id} - {tc_info['title']}", level=2)

            scenarios = grouped_executions[test_case_id]

            for scenario_name in sorted(scenarios.keys()):
                scenario_executions = scenarios[scenario_name]

                # Scenario subheading
                scenario_heading = doc.add_heading(f"Scenario: {scenario_name}", level=3)
                scenario_heading.paragraph_format.left_indent = Inches(0.25)

                # Scenario execution count
                doc.add_paragraph(f"({len(scenario_executions)} execution{'s' if len(scenario_executions) > 1 else ''})")

                # Create table for this scenario's executions
                table = doc.add_table(rows=1, cols=7)
                table.style = "Light Grid Accent 1"

                # Header row
                headers = ["Date", "Executed By", "Status", "Duration (min)", "Passed", "Failed", "Total Steps"]
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].runs[0].font.bold = True

                # Add executions for this scenario
                for exec_data in scenario_executions:
                    execution = exec_data['execution']
                    steps = exec_data['steps']

                    row_cells = table.add_row().cells
                    row_cells[0].text = execution.execution_date.strftime('%Y-%m-%d %H:%M')
                    row_cells[1].text = execution.executed_by
                    row_cells[2].text = execution.status

                    # Color code status
                    status_run = row_cells[2].paragraphs[0].runs[0]
                    if execution.status == 'PASSED':
                        status_run.font.color.rgb = RGBColor(0, 128, 0)
                    elif execution.status == 'FAILED':
                        status_run.font.color.rgb = RGBColor(255, 0, 0)
                    elif execution.status == 'BLOCKED':
                        status_run.font.color.rgb = RGBColor(255, 165, 0)
                    status_run.font.bold = True

                    row_cells[3].text = f"{execution.execution_time_minutes:.1f}" if execution.execution_time_minutes else "0.0"

                    # Calculate scenario-specific step counts
                    if steps:
                        scenario_passed = sum(1 for s in steps if s.get('status') == 'PASSED')
                        scenario_failed = sum(1 for s in steps if s.get('status') == 'FAILED')
                        scenario_total = len(steps)
                        row_cells[4].text = str(scenario_passed)
                        row_cells[5].text = str(scenario_failed)
                        row_cells[6].text = str(scenario_total)
                    else:
                        row_cells[4].text = str(execution.passed_steps)
                        row_cells[5].text = str(execution.failed_steps)
                        row_cells[6].text = str(execution.total_steps)

                doc.add_paragraph()

    def _calculate_project_metrics(self, projects: List[ProjectDB]) -> List[Dict]:
        """Calculate metrics for each project"""
        project_metrics = []

        for project in projects:
            # Get counts for this project
            stories_count = self.db.query(UserStoryDB).filter(UserStoryDB.project_id == project.id).count()
            test_cases_count = self.db.query(TestCaseDB).filter(TestCaseDB.project_id == project.id).count()
            bugs_count = self.db.query(BugReportDB).filter(BugReportDB.project_id == project.id).count()

            # Calculate test coverage
            stories_with_tests = self.db.query(UserStoryDB).filter(
                UserStoryDB.project_id == project.id,
                UserStoryDB.id.in_(
                    self.db.query(TestCaseDB.user_story_id).filter(TestCaseDB.project_id == project.id)
                )
            ).count()

            coverage = (stories_with_tests / stories_count * 100) if stories_count > 0 else 0

            # Calculate health score
            coverage_score = (coverage / 100) * 40
            bug_score = max(0, (1 - (bugs_count / (stories_count or 1))) * 30)
            test_score = max(0, ((test_cases_count / (stories_count or 1)) / 3) * 30)
            health_score = min(100, coverage_score + bug_score + test_score)

            # Determine risk level
            risk_factors = []
            if coverage < 50:
                risk_factors.append('baja cobertura')
            if bugs_count > stories_count * 0.3:
                risk_factors.append('alto número de bugs')
            if test_cases_count < stories_count:
                risk_factors.append('pocos test cases')

            if len(risk_factors) >= 2:
                risk_level = 'ALTO'
            elif len(risk_factors) == 1:
                risk_level = 'MEDIO'
            else:
                risk_level = 'BAJO'

            project_metrics.append({
                'name': project.name,
                'id': project.id,
                'status': project.status,
                'stories': stories_count,
                'test_cases': test_cases_count,
                'bugs': bugs_count,
                'coverage': coverage,
                'health_score': health_score,
                'risk_level': risk_level,
                'risk_factors': risk_factors
            })

        return project_metrics

    def _create_consolidated_document(
        self,
        projects: List[ProjectDB],
        project_metrics: List[Dict],
        projects_by_health: List[Dict],
        at_risk_projects: List[Dict],
        total_stories: int,
        total_test_cases: int,
        total_bugs: int,
        avg_coverage: float,
        avg_health: float
    ) -> Document:
        """Create Word document for consolidated report"""
        doc = Document()

        # Title
        title = doc.add_heading("Reporte Consolidado de Proyectos", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Header info
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"Quality Mission Control System\n").bold = True
        info_para.add_run(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
        info_para.add_run(f"Total de Proyectos: {len(projects)}")

        # Executive Summary
        doc.add_heading("Resumen Ejecutivo", level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Proyectos Totales: {len(projects)}\n")
        summary_para.add_run(f"User Stories Totales: {total_stories}\n")
        summary_para.add_run(f"Test Cases Totales: {total_test_cases}\n")
        summary_para.add_run(f"Bugs Totales: {total_bugs}\n")
        summary_para.add_run(f"Cobertura Promedio: {avg_coverage:.1f}%\n")
        summary_para.add_run(f"Health Score Promedio: {avg_health:.1f}/100")

        # At-Risk Projects
        self._add_risk_section_to_document(doc, at_risk_projects)

        # Detailed Metrics Table
        self._add_metrics_table_to_document(doc, projects_by_health)

        # Top Performers
        self._add_top_performers_to_document(doc, projects_by_health)

        # Recommendations
        self._add_recommendations_to_document(doc, project_metrics, avg_coverage, total_bugs, total_stories)

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph("Reporte generado automáticamente por Quality Mission Control System")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        return doc

    def _add_risk_section_to_document(self, doc: Document, at_risk_projects: List[Dict]):
        """Add risk section to document"""
        doc.add_heading("Proyectos que Requieren Atención", level=1)

        if at_risk_projects:
            risk_para = doc.add_paragraph()
            risk_para.add_run(f"{len(at_risk_projects)} proyecto(s) identificado(s) con riesgo medio o alto.\n\n").bold = True
            risk_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)

            for proj in at_risk_projects:
                risk_item = doc.add_paragraph(style='List Bullet')
                risk_run = risk_item.add_run(f"{proj['name']} - Riesgo {proj['risk_level']}")
                if proj['risk_level'] == 'ALTO':
                    risk_run.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    risk_run.font.color.rgb = RGBColor(255, 165, 0)
                risk_run.font.bold = True

                if proj['risk_factors']:
                    risk_item.add_run(f"\n  Factores: {', '.join(proj['risk_factors'])}")
        else:
            doc.add_paragraph("No hay proyectos en riesgo. ¡Excelente trabajo!")

    def _add_metrics_table_to_document(self, doc: Document, projects_by_health: List[Dict]):
        """Add detailed metrics table to document"""
        doc.add_heading("Métricas por Proyecto", level=1)

        table = doc.add_table(rows=1, cols=7)
        table.style = "Light Grid Accent 1"

        # Header row
        headers = ["Proyecto", "Stories", "Tests", "Bugs", "Cobertura", "Health Score", "Riesgo"]
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add projects (sorted by health score)
        for proj in projects_by_health:
            row_cells = table.add_row().cells
            row_cells[0].text = proj['name']
            row_cells[1].text = str(proj['stories'])
            row_cells[2].text = str(proj['test_cases'])
            row_cells[3].text = str(proj['bugs'])

            # Coverage with color coding
            coverage_cell = row_cells[4]
            coverage_cell.text = f"{proj['coverage']:.1f}%"
            coverage_run = coverage_cell.paragraphs[0].runs[0]
            if proj['coverage'] >= 70:
                coverage_run.font.color.rgb = RGBColor(0, 128, 0)
            elif proj['coverage'] >= 50:
                coverage_run.font.color.rgb = RGBColor(255, 165, 0)
            else:
                coverage_run.font.color.rgb = RGBColor(255, 0, 0)

            # Health Score with color coding
            health_cell = row_cells[5]
            health_cell.text = f"{proj['health_score']:.0f}/100"
            health_run = health_cell.paragraphs[0].runs[0]
            if proj['health_score'] >= 70:
                health_run.font.color.rgb = RGBColor(0, 128, 0)
            elif proj['health_score'] >= 50:
                health_run.font.color.rgb = RGBColor(255, 165, 0)
            else:
                health_run.font.color.rgb = RGBColor(255, 0, 0)

            # Risk Level with color coding
            risk_cell = row_cells[6]
            risk_cell.text = proj['risk_level']
            risk_run = risk_cell.paragraphs[0].runs[0]
            risk_run.font.bold = True
            if proj['risk_level'] == 'ALTO':
                risk_run.font.color.rgb = RGBColor(255, 0, 0)
            elif proj['risk_level'] == 'MEDIO':
                risk_run.font.color.rgb = RGBColor(255, 165, 0)
            else:
                risk_run.font.color.rgb = RGBColor(0, 128, 0)

            # Center align numeric columns
            for i in range(1, 7):
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_top_performers_to_document(self, doc: Document, projects_by_health: List[Dict]):
        """Add top performers section to document"""
        doc.add_heading("Proyectos con Mejor Desempeño", level=1)

        top_projects = projects_by_health[:3]
        top_para = doc.add_paragraph()
        top_para.add_run(f"Top {len(top_projects)} proyectos por Health Score:\n\n").bold = True
        top_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)

        for i, proj in enumerate(top_projects, 1):
            top_item = doc.add_paragraph(style='List Number')
            top_run = top_item.add_run(f"{proj['name']} - {proj['health_score']:.0f}/100")
            top_run.font.bold = True
            top_run.font.color.rgb = RGBColor(0, 128, 0)
            top_item.add_run(f"\n  Cobertura: {proj['coverage']:.1f}% | Tests: {proj['test_cases']} | Bugs: {proj['bugs']}")

    def _add_recommendations_to_document(
        self,
        doc: Document,
        project_metrics: List[Dict],
        avg_coverage: float,
        total_bugs: int,
        total_stories: int
    ):
        """Add recommendations section to document"""
        doc.add_heading("Recomendaciones", level=1)

        recommendations = []

        if avg_coverage < 60:
            recommendations.append("Incrementar la cobertura de tests en todos los proyectos. Objetivo: 70%+")

        if total_bugs > total_stories * 0.2:
            recommendations.append("Reducir la cantidad de bugs. Considerar sesiones de revisión de código y testing más exhaustivo.")

        low_coverage_projects = [p['name'] for p in project_metrics if p['coverage'] < 50]
        if low_coverage_projects:
            recommendations.append(f"Priorizar cobertura de tests en: {', '.join(low_coverage_projects)}")

        high_bug_projects = [p['name'] for p in project_metrics if p['bugs'] > p['stories'] * 0.3]
        if high_bug_projects:
            recommendations.append(f"Enfocar esfuerzos de QA en: {', '.join(high_bug_projects)}")

        if not recommendations:
            recommendations.append("Excelente trabajo. Todos los proyectos están en buen estado.")

        for rec in recommendations:
            rec_para = doc.add_paragraph(rec, style='List Bullet')
            if "Excelente" in rec:
                rec_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)


def get_report_service(db: Session) -> ReportService:
    """Dependency injection helper for FastAPI"""
    return ReportService(db)
