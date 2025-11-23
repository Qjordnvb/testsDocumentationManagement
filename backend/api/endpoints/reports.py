from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from pathlib import Path
from datetime import datetime
import os

from backend.database import get_db, ProjectDB, UserStoryDB, TestCaseDB, BugReportDB, TestExecutionDB
from backend.models import UserStory, TestCase, BugReport, BugSeverity, BugPriority, BugStatus, BugType
from backend.generators import TestPlanGenerator
from backend.generators.bug_report_generator import BugReportGenerator
from backend.config import settings

router = APIRouter()

@router.post("/generate-test-plan")
async def generate_test_plan(
    project_id: str = Query(..., description="Project ID to generate test plan for"),
    format: str = Query(default="both", description="Format: pdf, docx, or both"),
    db: Session = Depends(get_db)
):
    """
    Generate test plan document for a specific project
    """
    # Validate project exists
    project = db.query(ProjectDB).filter(ProjectDB.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail=f"Project {project_id} not found")

    # Get user stories and test cases for this project only
    user_stories_db = db.query(UserStoryDB).filter(UserStoryDB.project_id == project_id).all()
    test_cases_db = db.query(TestCaseDB).filter(TestCaseDB.project_id == project_id).all()

    # Convert to models (simplified)
    user_stories = [
        UserStory(
            id=s.id,
            title=s.title,
            description=s.description,
            priority=s.priority,
            status=s.status
        )
        for s in user_stories_db
    ]

    test_cases = [
        TestCase(
            id=tc.id,
            title=tc.title,
            description=tc.description,
            user_story_id=tc.user_story_id,
            test_type=tc.test_type,
            priority=tc.priority,
            status=tc.status
        )
        for tc in test_cases_db
    ]

    # Generate test plan
    settings.ensure_directories()
    test_plan_gen = TestPlanGenerator()
    files = test_plan_gen.generate_test_plan(
        user_stories=user_stories,
        test_cases=test_cases,
        output_dir=settings.output_dir,
        project_name=project.name,  # Use project name from database
        format=format
    )

    return {
        "message": "Test plan generated successfully",
        "files": files
    }

@router.get("/download/{filename}")
async def download_file(filename: str):
    """Download generated file"""
    file_path = Path(settings.output_dir) / filename

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/octet-stream"
    )


@router.get("/projects/{project_id}/reports/bug-summary")
async def generate_bug_summary_report(
    project_id: str,
    db: Session = Depends(get_db)
):
    """
    Generate Bug Summary Report for Dev Team
    Returns a Word document with all bugs for the project
    """
    # Validate project exists
    project = db.query(ProjectDB).filter(ProjectDB.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail=f"Project {project_id} not found")

    # Get all bugs for the project
    bugs_db = db.query(BugReportDB).filter(BugReportDB.project_id == project_id).all()

    if not bugs_db:
        raise HTTPException(status_code=404, detail="No bugs found for this project")

    # Convert to Pydantic models
    import json
    bugs = []
    for bug_db in bugs_db:
        bug = BugReport(
            id=bug_db.id,
            title=bug_db.title,
            description=bug_db.description,
            steps_to_reproduce=bug_db.steps_to_reproduce.split('\n') if bug_db.steps_to_reproduce else [],
            expected_behavior=bug_db.expected_behavior or "",
            actual_behavior=bug_db.actual_behavior or "",
            severity=BugSeverity(bug_db.severity),
            priority=BugPriority(bug_db.priority),
            bug_type=BugType(bug_db.bug_type),
            status=BugStatus(bug_db.status),
            environment=bug_db.environment,
            browser=bug_db.browser,
            os=bug_db.os,
            version=bug_db.version,
            user_story_id=bug_db.user_story_id,
            test_case_id=bug_db.test_case_id,
            scenario_name=bug_db.scenario_name,
            screenshots=json.loads(bug_db.screenshots) if bug_db.screenshots else [],
            logs=bug_db.logs,
            notes=bug_db.notes,
            workaround=bug_db.workaround,
            root_cause=bug_db.root_cause,
            fix_description=bug_db.fix_description,
            reported_by=bug_db.reported_by or "Unknown",
            assigned_to=bug_db.assigned_to,
            verified_by=bug_db.verified_by,
            reported_date=bug_db.reported_date,
            assigned_date=bug_db.assigned_date,
            fixed_date=bug_db.fixed_date,
            verified_date=bug_db.verified_date,
            closed_date=bug_db.closed_date
        )
        bugs.append(bug)

    # Generate report
    generator = BugReportGenerator()
    output_dir = "output/reports"
    filename = f"BugSummary_{project.name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    file_path = generator.generate_bulk_report(bugs, output_dir, filename)

    # Return file
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )


@router.get("/projects/{project_id}/reports/test-execution-summary")
async def generate_test_execution_report(
    project_id: str,
    db: Session = Depends(get_db)
):
    """
    Generate Test Execution Summary Report for QA Manager
    Returns a Word document with execution statistics and details grouped by Test Case and Scenario
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import json
    from collections import defaultdict

    # Validate project exists
    project = db.query(ProjectDB).filter(ProjectDB.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail=f"Project {project_id} not found")

    # Get test cases for the project
    test_cases = db.query(TestCaseDB).filter(TestCaseDB.project_id == project_id).all()

    if not test_cases:
        raise HTTPException(status_code=404, detail="No test cases found for this project")

    # Get all executions for these test cases
    test_case_ids = [tc.id for tc in test_cases]
    executions = db.query(TestExecutionDB).filter(
        TestExecutionDB.test_case_id.in_(test_case_ids)
    ).order_by(TestExecutionDB.execution_date.desc()).all()

    # Calculate statistics
    total_tests = len(test_cases)
    total_executions = len(executions)

    # Execution status counts
    status_counts = {
        'PASSED': 0,
        'FAILED': 0,
        'BLOCKED': 0,
        'SKIPPED': 0,
        'NOT_RUN': 0
    }

    for execution in executions:
        status_counts[execution.status] = status_counts.get(execution.status, 0) + 1

    # Test case status (latest execution)
    test_case_latest_status = {}
    test_case_info = {}  # Store test case info for later use
    for tc in test_cases:
        test_case_info[tc.id] = {
            'title': tc.title,
            'type': tc.test_type,
            'priority': tc.priority
        }

        # Get latest execution for this test case
        latest = db.query(TestExecutionDB).filter(
            TestExecutionDB.test_case_id == tc.id
        ).order_by(TestExecutionDB.execution_date.desc()).first()

        if latest:
            test_case_latest_status[tc.id] = latest.status
        else:
            test_case_latest_status[tc.id] = 'NOT_RUN'

    # Group executions by Test Case and Scenario
    grouped_executions = defaultdict(lambda: defaultdict(list))

    for execution in executions:
        test_case_id = execution.test_case_id

        # Parse step_results to extract scenario information
        if execution.step_results:
            try:
                step_results = json.loads(execution.step_results)

                # Group steps by scenario
                scenario_steps = defaultdict(list)
                for step in step_results:
                    scenario_name = step.get('scenario', 'Default Scenario')
                    scenario_steps[scenario_name].append(step)

                # Create execution record for each scenario
                for scenario_name, steps in scenario_steps.items():
                    grouped_executions[test_case_id][scenario_name].append({
                        'execution': execution,
                        'steps': steps
                    })
            except json.JSONDecodeError:
                # Fallback if JSON parsing fails
                grouped_executions[test_case_id]['Default Scenario'].append({
                    'execution': execution,
                    'steps': []
                })
        else:
            # No step results, add to default scenario
            grouped_executions[test_case_id]['Default Scenario'].append({
                'execution': execution,
                'steps': []
            })

    # Create document
    doc = Document()

    # Title
    title = doc.add_heading(f"Test Execution Summary Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f"Project: {project.name}\n").bold = True
    info_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    info_para.add_run(f"Total Test Cases: {total_tests}\n")
    info_para.add_run(f"Total Executions: {total_executions}")

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(f"Total Test Cases: {total_tests}")
    doc.add_paragraph(f"Total Test Executions: {total_executions}")

    # Execution Status Distribution
    doc.add_heading("Execution Status Distribution", level=2)
    status_table = doc.add_table(rows=1, cols=2)
    status_table.style = "Light Grid Accent 1"

    # Header
    header_cells = status_table.rows[0].cells
    header_cells[0].text = "Status"
    header_cells[1].text = "Count"
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    # Add status counts
    for status, count in sorted(status_counts.items()):
        row_cells = status_table.add_row().cells
        row_cells[0].text = status
        row_cells[1].text = str(count)

    # Pass Rate
    if total_executions > 0:
        pass_rate = (status_counts['PASSED'] / total_executions) * 100
        doc.add_paragraph()
        pass_rate_para = doc.add_paragraph()
        pass_rate_para.add_run(f"Overall Pass Rate: {pass_rate:.1f}%").bold = True
        if pass_rate >= 80:
            pass_rate_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif pass_rate >= 60:
            pass_rate_para.runs[0].font.color.rgb = RGBColor(255, 165, 0)  # Orange
        else:
            pass_rate_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red

    # Detailed Execution Results (Grouped by Test Case and Scenario)
    doc.add_heading("Execution Results by Test Case and Scenario", level=1)

    if not grouped_executions:
        doc.add_paragraph("No execution results to display.")
    else:
        for test_case_id in sorted(grouped_executions.keys()):
            # Test Case heading
            tc_info = test_case_info.get(test_case_id, {'title': 'Unknown', 'type': 'N/A', 'priority': 'N/A'})
            doc.add_heading(f"Test Case: {test_case_id} - {tc_info['title']}", level=2)

            scenarios = grouped_executions[test_case_id]

            for scenario_name in sorted(scenarios.keys()):
                scenario_executions = scenarios[scenario_name]

                # Scenario subheading
                scenario_heading = doc.add_heading(f"Scenario: {scenario_name}", level=3)
                scenario_heading.paragraph_format.left_indent = Inches(0.25)

                # Scenario execution count
                doc.add_paragraph(f"({len(scenario_executions)} execution{'s' if len(scenario_executions) > 1 else ''})")

                # Create table for this scenario's executions
                table = doc.add_table(rows=1, cols=7)
                table.style = "Light Grid Accent 1"

                # Header row
                headers = ["Date", "Executed By", "Status", "Duration (min)", "Passed", "Failed", "Total Steps"]
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].runs[0].font.bold = True

                # Add executions for this scenario
                for exec_data in scenario_executions:
                    execution = exec_data['execution']
                    steps = exec_data['steps']

                    row_cells = table.add_row().cells
                    row_cells[0].text = execution.execution_date.strftime('%Y-%m-%d %H:%M')
                    row_cells[1].text = execution.executed_by
                    row_cells[2].text = execution.status

                    # Color code status
                    status_run = row_cells[2].paragraphs[0].runs[0]
                    if execution.status == 'PASSED':
                        status_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                    elif execution.status == 'FAILED':
                        status_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                    elif execution.status == 'BLOCKED':
                        status_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
                    status_run.font.bold = True

                    row_cells[3].text = f"{execution.execution_time_minutes:.1f}" if execution.execution_time_minutes else "0.0"

                    # Calculate scenario-specific step counts
                    if steps:
                        scenario_passed = sum(1 for s in steps if s.get('status') == 'PASSED')
                        scenario_failed = sum(1 for s in steps if s.get('status') == 'FAILED')
                        scenario_total = len(steps)
                        row_cells[4].text = str(scenario_passed)
                        row_cells[5].text = str(scenario_failed)
                        row_cells[6].text = str(scenario_total)
                    else:
                        row_cells[4].text = str(execution.passed_steps)
                        row_cells[5].text = str(execution.failed_steps)
                        row_cells[6].text = str(execution.total_steps)

                doc.add_paragraph()  # Add spacing between scenarios

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("Generated by Quality Mission Control System")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Save document
    output_dir = Path("output/reports")
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"TestExecution_{project.name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = output_dir / filename

    doc.save(str(file_path))

    # Return file
    return FileResponse(
        path=str(file_path),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )


@router.get("/reports/consolidated")
async def generate_consolidated_report(
    db: Session = Depends(get_db)
):
    """
    Generate Consolidated Report for Manager
    Returns a Word document with metrics from all projects
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Get all projects
    projects = db.query(ProjectDB).all()

    if not projects:
        raise HTTPException(status_code=404, detail="No projects found")

    # Collect metrics for each project
    project_metrics = []
    total_stories = 0
    total_test_cases = 0
    total_bugs = 0
    total_coverage = 0

    for project in projects:
        # Get counts for this project
        stories_count = db.query(UserStoryDB).filter(UserStoryDB.project_id == project.id).count()
        test_cases_count = db.query(TestCaseDB).filter(TestCaseDB.project_id == project.id).count()
        bugs_count = db.query(BugReportDB).filter(BugReportDB.project_id == project.id).count()

        # Calculate test coverage (stories with tests)
        stories_with_tests = db.query(UserStoryDB).filter(
            UserStoryDB.project_id == project.id,
            UserStoryDB.id.in_(
                db.query(TestCaseDB.user_story_id).filter(TestCaseDB.project_id == project.id)
            )
        ).count()

        coverage = (stories_with_tests / stories_count * 100) if stories_count > 0 else 0

        # Calculate health score (same formula as frontend)
        coverage_score = (coverage / 100) * 40
        bug_score = max(0, (1 - (bugs_count / (stories_count or 1))) * 30)
        test_score = max(0, ((test_cases_count / (stories_count or 1)) / 3) * 30)
        health_score = min(100, coverage_score + bug_score + test_score)

        # Determine risk level
        risk_factors = []
        if coverage < 50:
            risk_factors.append('baja cobertura')
        if bugs_count > stories_count * 0.3:
            risk_factors.append('alto número de bugs')
        if test_cases_count < stories_count:
            risk_factors.append('pocos test cases')

        if len(risk_factors) >= 2:
            risk_level = 'ALTO'
        elif len(risk_factors) == 1:
            risk_level = 'MEDIO'
        else:
            risk_level = 'BAJO'

        project_metrics.append({
            'name': project.name,
            'id': project.id,
            'status': project.status,
            'stories': stories_count,
            'test_cases': test_cases_count,
            'bugs': bugs_count,
            'coverage': coverage,
            'health_score': health_score,
            'risk_level': risk_level,
            'risk_factors': risk_factors
        })

        # Accumulate totals
        total_stories += stories_count
        total_test_cases += test_cases_count
        total_bugs += bugs_count
        total_coverage += coverage

    # Calculate averages
    avg_coverage = total_coverage / len(projects) if projects else 0
    avg_health = sum(p['health_score'] for p in project_metrics) / len(projects) if projects else 0

    # Sort projects by health score
    projects_by_health = sorted(project_metrics, key=lambda p: p['health_score'], reverse=True)
    at_risk_projects = [p for p in project_metrics if p['risk_level'] in ['ALTO', 'MEDIO']]

    # Create document
    doc = Document()

    # Title
    title = doc.add_heading("Reporte Consolidado de Proyectos", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f"Quality Mission Control System\n").bold = True
    info_para.add_run(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    info_para.add_run(f"Total de Proyectos: {len(projects)}")

    # Executive Summary
    doc.add_heading("Resumen Ejecutivo", level=1)

    summary_para = doc.add_paragraph()
    summary_para.add_run(f"Proyectos Totales: {len(projects)}\n")
    summary_para.add_run(f"User Stories Totales: {total_stories}\n")
    summary_para.add_run(f"Test Cases Totales: {total_test_cases}\n")
    summary_para.add_run(f"Bugs Totales: {total_bugs}\n")
    summary_para.add_run(f"Cobertura Promedio: {avg_coverage:.1f}%\n")
    summary_para.add_run(f"Health Score Promedio: {avg_health:.1f}/100")

    # At-Risk Projects
    doc.add_heading("Proyectos que Requieren Atención", level=1)

    if at_risk_projects:
        risk_para = doc.add_paragraph()
        risk_para.add_run(f"{len(at_risk_projects)} proyecto(s) identificado(s) con riesgo medio o alto.\n\n").bold = True
        risk_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)

        for proj in at_risk_projects:
            risk_item = doc.add_paragraph(style='List Bullet')
            risk_run = risk_item.add_run(f"{proj['name']} - Riesgo {proj['risk_level']}")
            if proj['risk_level'] == 'ALTO':
                risk_run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                risk_run.font.color.rgb = RGBColor(255, 165, 0)
            risk_run.font.bold = True

            if proj['risk_factors']:
                risk_item.add_run(f"\n  Factores: {', '.join(proj['risk_factors'])}")
    else:
        doc.add_paragraph("No hay proyectos en riesgo. ¡Excelente trabajo!")

    # Detailed Metrics Table
    doc.add_heading("Métricas por Proyecto", level=1)

    # Create table
    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"

    # Header row
    headers = ["Proyecto", "Stories", "Tests", "Bugs", "Cobertura", "Health Score", "Riesgo"]
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add projects (sorted by health score)
    for proj in projects_by_health:
        row_cells = table.add_row().cells
        row_cells[0].text = proj['name']
        row_cells[1].text = str(proj['stories'])
        row_cells[2].text = str(proj['test_cases'])
        row_cells[3].text = str(proj['bugs'])

        # Coverage with color coding
        coverage_cell = row_cells[4]
        coverage_cell.text = f"{proj['coverage']:.1f}%"
        coverage_run = coverage_cell.paragraphs[0].runs[0]
        if proj['coverage'] >= 70:
            coverage_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif proj['coverage'] >= 50:
            coverage_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        else:
            coverage_run.font.color.rgb = RGBColor(255, 0, 0)  # Red

        # Health Score with color coding
        health_cell = row_cells[5]
        health_cell.text = f"{proj['health_score']:.0f}/100"
        health_run = health_cell.paragraphs[0].runs[0]
        if proj['health_score'] >= 70:
            health_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif proj['health_score'] >= 50:
            health_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        else:
            health_run.font.color.rgb = RGBColor(255, 0, 0)  # Red

        # Risk Level with color coding
        risk_cell = row_cells[6]
        risk_cell.text = proj['risk_level']
        risk_run = risk_cell.paragraphs[0].runs[0]
        risk_run.font.bold = True
        if proj['risk_level'] == 'ALTO':
            risk_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        elif proj['risk_level'] == 'MEDIO':
            risk_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        else:
            risk_run.font.color.rgb = RGBColor(0, 128, 0)  # Green

        # Center align numeric columns
        for i in range(1, 7):
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Top Performers
    doc.add_heading("Proyectos con Mejor Desempeño", level=1)

    top_projects = projects_by_health[:3]  # Top 3
    top_para = doc.add_paragraph()
    top_para.add_run(f"Top {len(top_projects)} proyectos por Health Score:\n\n").bold = True
    top_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    for i, proj in enumerate(top_projects, 1):
        top_item = doc.add_paragraph(style='List Number')
        top_run = top_item.add_run(f"{proj['name']} - {proj['health_score']:.0f}/100")
        top_run.font.bold = True
        top_run.font.color.rgb = RGBColor(0, 128, 0)
        top_item.add_run(f"\n  Cobertura: {proj['coverage']:.1f}% | Tests: {proj['test_cases']} | Bugs: {proj['bugs']}")

    # Recommendations
    doc.add_heading("Recomendaciones", level=1)

    recommendations = []

    if avg_coverage < 60:
        recommendations.append("Incrementar la cobertura de tests en todos los proyectos. Objetivo: 70%+")

    if total_bugs > total_stories * 0.2:
        recommendations.append("Reducir la cantidad de bugs. Considerar sesiones de revisión de código y testing más exhaustivo.")

    low_coverage_projects = [p['name'] for p in project_metrics if p['coverage'] < 50]
    if low_coverage_projects:
        recommendations.append(f"Priorizar cobertura de tests en: {', '.join(low_coverage_projects)}")

    high_bug_projects = [p['name'] for p in project_metrics if p['bugs'] > p['stories'] * 0.3]
    if high_bug_projects:
        recommendations.append(f"Enfocar esfuerzos de QA en: {', '.join(high_bug_projects)}")

    if not recommendations:
        recommendations.append("Excelente trabajo. Todos los proyectos están en buen estado.")

    for rec in recommendations:
        rec_para = doc.add_paragraph(rec, style='List Bullet')
        if "Excelente" in rec:
            rec_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("Reporte generado automáticamente por Quality Mission Control System")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Save document
    output_dir = Path("output/reports")
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"Consolidated_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = output_dir / filename

    doc.save(str(file_path))

    # Return file
    return FileResponse(
        path=str(file_path),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )
